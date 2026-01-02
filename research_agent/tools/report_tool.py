"""Report generation tool for the Investment Research Agent.

This tool generates professional equity research reports in both Word (.docx)
and HTML formats. The HTML format uses OneNote-compatible elements for future
integration compatibility.
"""

import json
from datetime import datetime
from pathlib import Path
from typing import Any

from research_agent.tools.registry import registered_tool
from research_agent.tools.report_html import create_html_report

# Lazy import to avoid dependency issues if python-docx not installed
docx = None


def _ensure_docx():
    """Lazy import python-docx."""
    global docx
    if docx is None:
        try:
            import docx as _docx
            docx = _docx
        except ImportError:
            raise ImportError(
                "python-docx is required for report generation. "
                "Install it with: pip install python-docx"
            )
    return docx


def _get_outputs_dir() -> Path:
    """Get the outputs directory path."""
    from research_agent.config import get_outputs_dir
    return get_outputs_dir()


def _create_docx_report(
    company_name: str,
    ticker: str,
    sections: dict[str, Any],
) -> Path:
    """Create a Word document with the research report.

    Args:
        company_name: Company name (e.g., "DocuSign")
        ticker: Stock ticker (e.g., "DOCU")
        sections: Dict with report sections:
            - executive_summary: str
            - key_highlights: list[str]
            - financial_analysis: str
            - business_analysis: str
            - valuation: str
            - risks: list[dict] with 'title' and 'description'
            - key_takeaways: list[str]

    Returns:
        Path to the generated DOCX file.
    """
    _ensure_docx()
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    doc = docx.Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading(f"{company_name} ({ticker})", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Equity Research Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Report date
    date_para = doc.add_paragraph()
    date_para.add_run("Report Date: ").bold = True
    date_para.add_run(datetime.now().strftime("%B %d, %Y"))

    doc.add_paragraph()  # Spacing

    # Executive Summary
    if sections.get("executive_summary"):
        doc.add_heading("EXECUTIVE SUMMARY", level=1)
        para = doc.add_paragraph(sections["executive_summary"])
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Key Highlights
    if sections.get("key_highlights"):
        doc.add_heading("KEY HIGHLIGHTS", level=1)
        for highlight in sections["key_highlights"]:
            # Ensure highlight is a clean string (not iterated char-by-char)
            if isinstance(highlight, str):
                # Clean up the text - remove extra whitespace and ensure single-line
                clean_text = ' '.join(highlight.split())
                if clean_text:
                    doc.add_paragraph(clean_text, style='List Bullet')

    # Financial Analysis
    if sections.get("financial_analysis"):
        doc.add_heading("FINANCIAL ANALYSIS", level=1)
        para = doc.add_paragraph(sections["financial_analysis"])
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Business Analysis
    if sections.get("business_analysis"):
        doc.add_heading("BUSINESS ANALYSIS", level=1)
        para = doc.add_paragraph(sections["business_analysis"])
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Valuation
    if sections.get("valuation"):
        doc.add_heading("VALUATION ANALYSIS", level=1)
        para = doc.add_paragraph(sections["valuation"])
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Risks
    if sections.get("risks"):
        doc.add_heading("KEY RISKS", level=1)
        for i, risk in enumerate(sections["risks"], 1):
            if isinstance(risk, dict):
                title = risk.get('title', 'Risk')
                description = risk.get("description", "")
                # Clean up text
                if isinstance(title, str):
                    title = ' '.join(title.split())
                if isinstance(description, str):
                    description = ' '.join(description.split())
                doc.add_heading(f"{i}. {title}", level=2)
                para = doc.add_paragraph(description)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif isinstance(risk, str):
                clean_text = ' '.join(risk.split())
                if clean_text:
                    doc.add_paragraph(f"{i}. {clean_text}", style='List Bullet')

    # Key Takeaways
    if sections.get("key_takeaways"):
        doc.add_heading("KEY TAKEAWAYS", level=1)
        for takeaway in sections["key_takeaways"]:
            # Ensure takeaway is a clean string (not iterated char-by-char)
            if isinstance(takeaway, str):
                clean_text = ' '.join(takeaway.split())
                if clean_text:
                    doc.add_paragraph(clean_text, style='List Bullet')

    # Disclaimer
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.add_run("DISCLAIMER: ").bold = True
    disclaimer.add_run(
        "This analysis is for informational purposes only and does not constitute "
        "investment advice. Always conduct independent due diligence before making "
        "investment decisions. Past performance is not indicative of future results."
    )
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Save document
    filename = f"{company_name.replace(' ', '_')}_{ticker}_Analysis.docx"
    output_path = _get_outputs_dir() / filename
    doc.save(str(output_path))

    return output_path


@registered_tool(
    name="generate_report",
    description=(
        "Generate a professional equity research report in both Word (.docx) and HTML formats. "
        "The HTML format is OneNote-compatible for future integration. "
        "Use this after analyzing a company to create a formal report the user can view and download."
    ),
    parameters={
        "company_name": "The company name (e.g., 'DocuSign')",
        "ticker": "The stock ticker (e.g., 'DOCU')",
        "sections": (
            "JSON string with report sections: executive_summary, key_highlights, "
            "financial_analysis, business_analysis, valuation, risks, key_takeaways"
        ),
    },
    parameter_types={
        "company_name": str,
        "ticker": str,
        "sections": str,
    },
)
async def generate_report(args: dict[str, Any]) -> dict[str, Any]:
    """Generate an equity research report in HTML and DOCX formats.

    Args:
        args: Dict with:
            - company_name: Company name (e.g., "DocuSign")
            - ticker: Stock ticker (e.g., "DOCU")
            - sections: JSON string with report sections:
                {
                    "executive_summary": "...",
                    "key_highlights": ["...", "..."],
                    "financial_analysis": "...",
                    "business_analysis": "...",
                    "valuation": "...",
                    "risks": [{"title": "...", "description": "..."}, ...],
                    "key_takeaways": ["...", "..."]
                }

    Returns:
        Tool result with success/failure message and paths to both files.
    """
    try:
        company_name = args.get("company_name", "")
        ticker = args.get("ticker", "")
        sections_json = args.get("sections", "{}")

        if not company_name or not ticker:
            return {
                "content": [{
                    "type": "text",
                    "text": "Error: company_name and ticker are required."
                }],
                "is_error": True,
            }

        # Parse sections JSON
        try:
            sections = json.loads(sections_json)
        except json.JSONDecodeError as e:
            return {
                "content": [{
                    "type": "text",
                    "text": f"Error parsing sections JSON: {e}"
                }],
                "is_error": True,
            }

        # Generate HTML report (always available)
        html_path = create_html_report(company_name, ticker, sections)

        # Try to generate DOCX report (may fail if python-docx not installed)
        docx_path = None
        docx_error = None
        try:
            docx_path = _create_docx_report(company_name, ticker, sections)
        except ImportError as e:
            docx_error = str(e)

        # Build response message
        if docx_path:
            return {
                "content": [{
                    "type": "text",
                    "text": (
                        f"Reports generated successfully!\n\n"
                        f"HTML (viewable): {html_path}\n"
                        f"DOCX (downloadable): {docx_path}"
                    )
                }]
            }
        else:
            return {
                "content": [{
                    "type": "text",
                    "text": (
                        f"HTML report generated successfully!\n\n"
                        f"Saved to: {html_path}\n\n"
                        f"Note: DOCX generation skipped - {docx_error}"
                    )
                }]
            }

    except Exception as e:
        return {
            "content": [{
                "type": "text",
                "text": f"Error generating report: {e}"
            }],
            "is_error": True,
        }
