"""Report Builder - Builder pattern for generating research reports.

This module implements the Builder pattern for creating DOCX reports,
making the code more maintainable and testable.
"""

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

# Lazy import python-docx
_docx = None


def _ensure_docx():
    """Lazy import python-docx."""
    global _docx
    if _docx is None:
        try:
            import docx
            _docx = docx
        except ImportError:
            raise ImportError(
                "python-docx is required for report generation. "
                "Install it with: pip install python-docx"
            )
    return _docx


@dataclass
class ReportSections:
    """Value object for report sections."""

    executive_summary: str = ""
    key_highlights: list[str] = field(default_factory=list)
    financial_analysis: str = ""
    business_analysis: str = ""
    valuation: str = ""
    risks: list[dict[str, str]] = field(default_factory=list)
    key_takeaways: list[str] = field(default_factory=list)

    @classmethod
    def from_dict(cls, data: dict[str, Any]) -> "ReportSections":
        """Create ReportSections from a dictionary."""
        return cls(
            executive_summary=data.get("executive_summary", ""),
            key_highlights=data.get("key_highlights", []),
            financial_analysis=data.get("financial_analysis", ""),
            business_analysis=data.get("business_analysis", ""),
            valuation=data.get("valuation", ""),
            risks=data.get("risks", []),
            key_takeaways=data.get("key_takeaways", []),
        )


class DocxReportBuilder:
    """Builder for creating DOCX equity research reports.

    Uses the Builder pattern to construct reports step by step,
    with each section as a separate method.
    """

    def __init__(self, company_name: str, ticker: str, sections: ReportSections):
        """Initialize the builder with report data.

        Args:
            company_name: Company name (e.g., "DocuSign")
            ticker: Stock ticker (e.g., "DOCU")
            sections: Report sections data
        """
        self.company_name = company_name
        self.ticker = ticker
        self.sections = sections
        self._doc = None
        self._align_justify = None

    def _setup(self) -> None:
        """Initialize the document and styles."""
        docx = _ensure_docx()
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        self._doc = docx.Document()
        self._align_justify = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Set up default style
        style = self._doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

    def _add_header(self) -> None:
        """Add title, subtitle, and date."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Title
        title = self._doc.add_heading(f"{self.company_name} ({self.ticker})", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = self._doc.add_paragraph("Equity Research Report")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Report date
        date_para = self._doc.add_paragraph()
        date_para.add_run("Report Date: ").bold = True
        date_para.add_run(datetime.now().strftime("%B %d, %Y"))

        self._doc.add_paragraph()  # Spacing

    def _add_executive_summary(self) -> None:
        """Add executive summary section."""
        if not self.sections.executive_summary:
            return
        self._doc.add_heading("EXECUTIVE SUMMARY", level=1)
        para = self._doc.add_paragraph(self.sections.executive_summary)
        para.alignment = self._align_justify

    def _add_key_highlights(self) -> None:
        """Add key highlights section."""
        if not self.sections.key_highlights:
            return
        self._doc.add_heading("KEY HIGHLIGHTS", level=1)
        for highlight in self.sections.key_highlights:
            self._add_bullet(highlight)

    def _add_financial_analysis(self) -> None:
        """Add financial analysis section."""
        if not self.sections.financial_analysis:
            return
        self._doc.add_heading("FINANCIAL ANALYSIS", level=1)
        para = self._doc.add_paragraph(self.sections.financial_analysis)
        para.alignment = self._align_justify

    def _add_business_analysis(self) -> None:
        """Add business analysis section."""
        if not self.sections.business_analysis:
            return
        self._doc.add_heading("BUSINESS ANALYSIS", level=1)
        para = self._doc.add_paragraph(self.sections.business_analysis)
        para.alignment = self._align_justify

    def _add_valuation(self) -> None:
        """Add valuation analysis section."""
        if not self.sections.valuation:
            return
        self._doc.add_heading("VALUATION ANALYSIS", level=1)
        para = self._doc.add_paragraph(self.sections.valuation)
        para.alignment = self._align_justify

    def _add_risks(self) -> None:
        """Add key risks section."""
        if not self.sections.risks:
            return
        self._doc.add_heading("KEY RISKS", level=1)
        for i, risk in enumerate(self.sections.risks, 1):
            if isinstance(risk, dict):
                title = self._clean_text(risk.get('title', 'Risk'))
                description = self._clean_text(risk.get('description', ''))
                self._doc.add_heading(f"{i}. {title}", level=2)
                para = self._doc.add_paragraph(description)
                para.alignment = self._align_justify
            elif isinstance(risk, str):
                clean_text = self._clean_text(risk)
                if clean_text:
                    self._doc.add_paragraph(f"{i}. {clean_text}", style='List Bullet')

    def _add_key_takeaways(self) -> None:
        """Add key takeaways section."""
        if not self.sections.key_takeaways:
            return
        self._doc.add_heading("KEY TAKEAWAYS", level=1)
        for takeaway in self.sections.key_takeaways:
            self._add_bullet(takeaway)

    def _add_disclaimer(self) -> None:
        """Add legal disclaimer."""
        self._doc.add_paragraph()
        disclaimer = self._doc.add_paragraph()
        disclaimer.add_run("DISCLAIMER: ").bold = True
        disclaimer.add_run(
            "This analysis is for informational purposes only and does not constitute "
            "investment advice. Always conduct independent due diligence before making "
            "investment decisions. Past performance is not indicative of future results."
        )
        disclaimer.alignment = self._align_justify

    def _add_bullet(self, text: str) -> None:
        """Add a bullet point with cleaned text."""
        if isinstance(text, str):
            clean = self._clean_text(text)
            if clean:
                self._doc.add_paragraph(clean, style='List Bullet')

    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean text by normalizing whitespace."""
        if not isinstance(text, str):
            return ""
        return ' '.join(text.split())

    def build(self, output_path: Path) -> Path:
        """Build the complete document and save it.

        Args:
            output_path: Path where the document will be saved

        Returns:
            Path to the saved document
        """
        self._setup()
        self._add_header()
        self._add_executive_summary()
        self._add_key_highlights()
        self._add_financial_analysis()
        self._add_business_analysis()
        self._add_valuation()
        self._add_risks()
        self._add_key_takeaways()
        self._add_disclaimer()

        self._doc.save(str(output_path))
        return output_path


def build_docx_report(
    company_name: str,
    ticker: str,
    sections: dict[str, Any],
    output_dir: Path,
) -> Path:
    """Convenience function to build a DOCX report.

    Args:
        company_name: Company name (e.g., "DocuSign")
        ticker: Stock ticker (e.g., "DOCU")
        sections: Dict with report sections
        output_dir: Directory where the report will be saved

    Returns:
        Path to the generated DOCX file
    """
    report_sections = ReportSections.from_dict(sections)
    filename = f"{company_name.replace(' ', '_')}_{ticker}_Analysis.docx"
    output_path = output_dir / filename

    builder = DocxReportBuilder(company_name, ticker, report_sections)
    return builder.build(output_path)
